from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class DocumentText:
    text: str
    source_path: str
    kind: str  # "docx" | "pdf" | "txt"


def read_docx(path: str) -> DocumentText:
    from docx import Document  # python-docx

    p = Path(path)
    doc = Document(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text or "")
    # tables (dažnai literatūra nebūna, bet jei yra – įtraukiam)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return DocumentText(text="\n".join(parts).strip(), source_path=str(p), kind="docx")


_BULLET_OR_NUM_RE = re.compile(r"^\s*(?:\[\d{1,4}\]|(?:\d{1,4})[\.\)]|[-\u2022])\s*")
_HEADING_RE = re.compile(r"^\s*(references|bibliography|literat[ūu]ra|literatura|šaltiniai|saltiniai)\s*$", re.IGNORECASE)
_AUTHOR_COMMA_START_RE = re.compile(
    r"^\s*[A-Z][A-Za-z'`\-]{1,40}\s*,\s*(?:[A-Z]\.|[A-Z][a-z]{1,30}|[A-Z]\.[A-Z]\.)"
)
_AUTHOR_YEAR_START_RE = re.compile(
    r"^\s*[A-Z][A-Za-z'`\-]{2,40}\.?\s*(?:\(\s*(?:19|20)\d{2}[a-z]?\s*\)|\(\s*n\.d\.\s*\))",
    re.IGNORECASE,
)


def _looks_like_reference_start(line: str) -> bool:
    s = (line or "").strip()
    if not s:
        return False
    if _BULLET_OR_NUM_RE.match(s) or _HEADING_RE.match(s):
        return True
    if _AUTHOR_COMMA_START_RE.match(s):
        return True
    if _AUTHOR_YEAR_START_RE.match(s):
        return True
    return False


def _normalize_pdf_text(raw_text: str) -> str:
    """
    Normalizuoja PDF ištrauktą tekstą į stabilesnį TXT:
    - suvienodina tarpus;
    - sulipdo eilučių lūžius sakinio viduryje;
    - palieka naują eilutę prie numeruotų/bullet įrašų ir heading'ų.
    """
    if not raw_text:
        return ""

    text = raw_text.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    lines = [ln.rstrip() for ln in text.split("\n")]

    out: list[str] = []
    i = 0
    while i < len(lines):
        ln = re.sub(r"\s+", " ", lines[i]).strip()
        if not ln:
            out.append("")
            i += 1
            continue

        # De-hyphenation tarp eilučių: "crypto-\ncurrency" -> "cryptocurrency"
        cur = ln
        while cur.endswith("-") and i + 1 < len(lines):
            nxt = re.sub(r"\s+", " ", lines[i + 1]).strip()
            if not nxt:
                break
            cur = cur[:-1] + nxt
            i += 1

        # Jei kita eilutė nėra aiškiai naujo šaltinio pradžia, sulipdom kaip tąsa.
        while i + 1 < len(lines):
            nxt_raw = lines[i + 1]
            nxt = re.sub(r"\s+", " ", nxt_raw).strip()
            if not nxt:
                break
            if _looks_like_reference_start(nxt):
                break
            if cur.endswith((".", "!", "?", ":", ";")):
                break
            cur = f"{cur} {nxt}"
            i += 1

        out.append(cur)
        i += 1

    # Sumažinam kelių tuščių eilučių triukšmą iki vienos.
    normalized = "\n".join(out)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def read_pdf(path: str) -> DocumentText:
    import fitz  # pymupdf

    p = Path(path)
    doc = fitz.open(str(p))
    parts: list[str] = []
    for page in doc:
        # sort=True padeda nuoseklesnei skaitymo tvarkai.
        parts.append(page.get_text("text", sort=True))
    raw = "\n".join(parts).strip()
    cleaned = _normalize_pdf_text(raw)
    return DocumentText(text=cleaned, source_path=str(p), kind="pdf")


def read_text(path: str) -> DocumentText:
    p = Path(path)
    return DocumentText(text=p.read_text(encoding="utf-8", errors="ignore"), source_path=str(p), kind="txt")


def read_any(path: str) -> DocumentText:
    p = Path(path)
    suf = p.suffix.lower()
    if suf == ".docx":
        return read_docx(str(p))
    if suf == ".pdf":
        return read_pdf(str(p))
    return read_text(str(p))

